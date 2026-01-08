import streamlit as st
import io
from datetime import datetime
from docx import Document
from transformers import pipeline

# -------------------------------------------------
# PAGE CONFIG
# -------------------------------------------------
st.set_page_config(
    page_title="GenAI SOW Agent",
    layout="wide",
    page_icon="📝"
)

# -------------------------------------------------
# LOAD LLM (STREAMLIT CLOUD SAFE)
# -------------------------------------------------
@st.cache_resource
def load_llm():
    return pipeline(
        "text2text-generation",
        model="google/flan-t5-base",
        device="cpu"
    )

llm = load_llm()

def call_llm(prompt):
    out = llm(prompt, max_new_tokens=300, temperature=0.3)
    return out[0]["generated_text"].strip()

# -------------------------------------------------
# SESSION STATE INITIALIZATION
# -------------------------------------------------
def empty_row():
    return {"Name": "", "Title": "", "Email": ""}

if "sow_data" not in st.session_state:
    st.session_state.sow_data = {
        "metadata": {
            "customer_name": "",
            "industry": "Retail",
            "solution_type": "Multi Agent Store Advisor",
            "other_solution": "",
            "business_problem": ""
        },
        "sections": {
            "objective": "",
            "dependencies": "",
            "assumptions": "",
            "success_demo": "",
            "success_results": "",
            "scope": "",
            "architecture": "",
            "commercials": ""
        },
        "tables": {
            "partner": [empty_row()],
            "customer": [empty_row()],
            "cloud": [empty_row()],
            "escalation": [empty_row()]
        }
    }

# -------------------------------------------------
# GENERATION ENGINE (FIXED)
# -------------------------------------------------
def generate_sow_content():
    data = st.session_state.sow_data
    meta = data["metadata"]
    sec = data["sections"]

    solution = meta["other_solution"] if meta["solution_type"] == "Other (Please specify)" else meta["solution_type"]

    sec["objective"] = call_llm(
        f"Write a formal SOW Objective for a {solution} Proof of Concept in the {meta['industry']} industry. Business problem: {meta['business_problem']}"
    )

    sec["dependencies"] = call_llm(
        f"List customer dependencies for a {solution} POC."
    )

    sec["assumptions"] = call_llm(
        f"List delivery assumptions for a {solution} POC."
    )

    sec["success_demo"] = call_llm(
        f"List demonstration success criteria for a {solution} POC."
    )

    sec["success_results"] = call_llm(
        f"List measurable business outcomes for a {solution} POC."
    )

    sec["scope"] = call_llm(
        f"Describe a phase-wise technical scope (Discovery, Build, Test, Demo) for a {solution} POC."
    )

    sec["architecture"] = call_llm(
        f"Describe a high-level GenAI architecture for {solution}."
    )

    sec["commercials"] = call_llm(
        f"Write a short commercial and resource estimation note for a {solution} POC."
    )

# -------------------------------------------------
# TABLE EDITOR
# -------------------------------------------------
def render_table(title, key):
    st.subheader(title)
    rows = st.session_state.sow_data["tables"][key]

    for i, row in enumerate(rows):
        c1, c2, c3 = st.columns(3)
        row["Name"] = c1.text_input("Name", row["Name"], key=f"{key}_n_{i}")
        row["Title"] = c2.text_input("Title", row["Title"], key=f"{key}_t_{i}")
        row["Email"] = c3.text_input("Email / Contact", row["Email"], key=f"{key}_e_{i}")

    if st.button(f"➕ Add {title} Row", key=f"add_{key}"):
        rows.append(empty_row())
        st.rerun()

# -------------------------------------------------
# DOCX EXPORT
# -------------------------------------------------
def add_docx_table(doc, title, rows):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "Title"
    hdr[2].text = "Email / Contact Info"

    for r in rows:
        if any(r.values()):
            cells = table.add_row().cells
            cells[0].text = r["Name"]
            cells[1].text = r["Title"]
            cells[2].text = r["Email"]

def create_docx(data):
    doc = Document()
    doc.add_heading("Statement of Work (SOW)", 0)

    meta = data["metadata"]
    sec = data["sections"]

    doc.add_paragraph(f"Customer: {meta['customer_name']}")
    doc.add_paragraph(f"Industry: {meta['industry']}")
    doc.add_paragraph(f"Solution: {meta['solution_type']}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")

    doc.add_heading("2.1 Objective", 1)
    doc.add_paragraph(sec["objective"])

    doc.add_heading("2.2 Stakeholders", 1)
    add_docx_table(doc, "Partner Executive Sponsor", data["tables"]["partner"])
    add_docx_table(doc, "Customer Executive Sponsor", data["tables"]["customer"])
    add_docx_table(doc, "Cloud Executive Sponsor", data["tables"]["cloud"])
    add_docx_table(doc, "Escalation Matrix", data["tables"]["escalation"])

    doc.add_heading("2.3 Assumptions & Dependencies", 1)
    doc.add_paragraph("Dependencies:\n" + sec["dependencies"])
    doc.add_paragraph("Assumptions:\n" + sec["assumptions"])

    doc.add_heading("2.4 Success Criteria", 1)
    doc.add_paragraph("Demonstration:\n" + sec["success_demo"])
    doc.add_paragraph("Results:\n" + sec["success_results"])

    doc.add_heading("3 Scope of Work", 1)
    doc.add_paragraph(sec["scope"])

    doc.add_heading("4 Solution Architecture", 1)
    doc.add_paragraph(sec["architecture"])

    doc.add_heading("5 Commercials", 1)
    doc.add_paragraph(sec["commercials"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# -------------------------------------------------
# UI
# -------------------------------------------------
def main():
    data = st.session_state.sow_data
    meta = data["metadata"]
    sec = data["sections"]

    st.sidebar.title("🛠️ Inputs")

    meta["customer_name"] = st.sidebar.text_input("Customer Name", meta["customer_name"])

    meta["industry"] = st.sidebar.selectbox(
        "Industry",
        [
            "Retail", "Financial Services", "Banking", "Insurance",
            "Healthcare", "Pharmaceuticals", "Manufacturing",
            "Automotive", "Telecom", "Energy", "Utilities",
            "Travel & Hospitality", "Logistics", "E-commerce",
            "Media & Entertainment", "Legal", "Education", "Public Sector"
        ]
    )

    solution_options = [
        "Multi Agent Store Advisor",
        "Intelligent Search",
        "Recommendation",
        "AI Agents Demand Forecasting",
        "Banner Audit using LLM",
        "Image Enhancement",
        "Virtual Try-On",
        "Agentic AI L1 Support",
        "Product Listing Standardization",
        "AI Agents Based Pricing Module",
        "Cost, Margin Visibility & Insights using LLM",
        "AI Trend Simulator",
        "Virtual Data Analyst (Text to SQL)",
        "Multilingual Call Analysis",
        "Customer Review Analysis",
        "Sales Co-Pilot",
        "Research Co-Pilot",
        "Product Copy Generator",
        "Multi-agent e-KYC & Onboarding",
        "Document / Report Audit",
        "RBI Circular Scraping & Insights Bot",
        "Visual Inspection",
        "AIoT based CCTV Surveillance",
        "Multilingual Voice Bot",
        "SOP Creation",
        "Other (Please specify)"
    ]

    meta["solution_type"] = st.sidebar.selectbox("Solution Type", solution_options)

    if meta["solution_type"] == "Other (Please specify)":
        meta["other_solution"] = st.sidebar.text_input("Specify Solution")

    meta["business_problem"] = st.sidebar.text_area("Business Problem")

    if st.sidebar.button("🪄 Generate SOW Content"):
        generate_sow_content()
        st.success("Content generated successfully")

    st.title("📄 Editable SOW")

    sec["objective"] = st.text_area("2.1 Objective", sec["objective"], height=150)
    render_table("Partner Executive Sponsor", "partner")
    render_table("Customer Executive Sponsor", "customer")
    render_table("Cloud Executive Sponsor", "cloud")
    render_table("Escalation Matrix", "escalation")

    sec["dependencies"] = st.text_area("Dependencies", sec["dependencies"], height=120)
    sec["assumptions"] = st.text_area("Assumptions", sec["assumptions"], height=120)
    sec["success_demo"] = st.text_area("Success Demonstration", sec["success_demo"], height=120)
    sec["success_results"] = st.text_area("Expected Results", sec["success_results"], height=120)
    sec["scope"] = st.text_area("Scope of Work", sec["scope"], height=160)
    sec["architecture"] = st.text_area("Architecture", sec["architecture"], height=140)
    sec["commercials"] = st.text_area("Commercials", sec["commercials"], height=120)

    docx = create_docx(data)
    st.download_button(
        "📥 Download SOW (DOCX)",
        docx,
        file_name="Generated_SOW.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

if __name__ == "__main__":
    main()
